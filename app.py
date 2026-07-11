import streamlit as st
import re
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="MCQ Formatter", layout="wide")

st.title("📚 Bhargav's MCQ Formatter (English & Assamese)")

raw_text = st.text_area(
    "Paste Questions Here (Supports Single-line & Multi-line Options)",
    height=400
)

if st.button("Format MCQs"):

    # Split blocks on question headers (Handles 'Q1.', 'প্ৰশ্ন ১.', 'প্ৰ. ১.')
    blocks = re.split(r'\n(?=(?:Q|প্ৰশ্ন|প্ৰ\.)\s*(?:\d+|[০-৯]+))', raw_text.strip())

    doc = Document()

    # Expanded answer map to translate both formats to numeric indices
    answer_map = {
        "A": "1", "B": "2", "C": "3", "D": "4",
        "ক": "1", "খ": "2", "গ": "3", "ঘ": "4"
    }

    total_questions = 0

    for block in blocks:

        if not block.strip():
            continue

        lines = [x.strip() for x in block.split("\n") if x.strip()]

        question_lines = []
        options = []
        solution_lines = []
        answer = ""

        after_options = False

        for line in lines:

            # 1. Parse Answer Lines (Handles "Answer: (b)", "উত্তৰ: (খ)", etc.)
            if line.lower().startswith("answer") or line.startswith("উত্তৰ"):
                ans_match = re.search(r'(?:answer|উত্তৰ)\s*:\s*\(?([a-dA-Dকখগঘ])[\.\)]?', line, re.IGNORECASE)
                if ans_match:
                    answer = ans_match.group(1).upper()
                    explanation_part = line[ans_match.end():].strip()
                    if explanation_part:
                        solution_lines.append(explanation_part)
                else:
                    answer_text = line.split(":", 1)[-1].strip()
                    if len(answer_text) == 1:
                        answer = answer_text.upper()
                    else:
                        solution_lines.append(answer_text)
                after_options = True

            # 2. Parse Option lines (Handles both Single-line and Inline Packed Options)
            else:
                # Find all option delimiters on this line (e.g., A), B), C), D) or (a), (b)...)
                opt_matches = list(re.finditer(r'(?:^|\s)(\(?[A-Da-dকখগঘ][\.\)])(?=\s+|$)', line))
                
                # Check if this line qualifies as an options line
                is_option_line = False
                if opt_matches:
                    first_marker = opt_matches[0].group(1).strip()
                    # It's an option line if it starts with a marker OR contains multiple markers inline
                    if line.strip().startswith(first_marker) or len(opt_matches) > 1:
                        is_option_line = True
                
                if is_option_line:
                    after_options = True
                    # Dynamically slice the single line into individual options
                    for i in range(len(opt_matches)):
                        start_pos = opt_matches[i].end()
                        end_pos = opt_matches[i+1].start() if i + 1 < len(opt_matches) else len(line)
                        opt_text = line[start_pos:end_pos].strip()
                        if opt_text:
                            options.append(opt_text)
                            
                # 3. Everything after options becomes solution text
                elif after_options:
                    solution_lines.append(line)

                # 4. Question text accumulation
                else:
                    question_lines.append(line)

        # Join question components
        question = "\n".join(question_lines)

        # Remove prefix pattern (Handles 'Q99.', 'প্ৰশ্ন ৯৯.', 'প্ৰ. ৯৯.')
        question = re.sub(r'^(?:Q|প্ৰশ্ন|প্ৰ\.)\s*(?:\d+|[০-৯]+)[\.\s]*', '', question)

        # Remove bracket prefix metadata if present
        question = re.sub(r'^\([^)]*\)\s*', '', question)

        # --- NEW ADDITION: Force inline statement markers (I., II., 1., ২., etc.) onto new lines ---
        question = re.sub(r'\s+(?=(?:[IVXivx]+|\d+|[০-৯]+)[\.\)]\s|\((?:[IVXivx]+|\d+|[০-৯]+)\)\s)', '\n', question)

        solution = "\n".join(solution_lines).strip()

        # Remove inline answers from solution if duplicated
        solution = re.sub(
            r'(?:Answer|উত্তৰ)\s*:\s*(?:[A-D]|[কখগঘ])',
            '',
            solution,
            flags=re.IGNORECASE
        ).strip()

        # Numeric translation mapping
        answer_numeric = answer_map.get(answer, answer)

        # Create table configuration structure
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        table.columns[0].width = Inches(1.3)
        table.columns[1].width = Inches(5.8)

        def add_row(label, value):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
            row[0].width = Inches(1.3)
            row[1].width = Inches(5.8)

        # Populate structured output rows
        add_row("Question", question)
        add_row("Type", "multiple_choice")

        for option in options:
            add_row("Option", option)

        add_row("Answer", answer_numeric)
        add_row("Solution", solution)
        add_row("Positive Marks", "1")
        add_row("Negative Marks", "0.25")

        doc.add_paragraph()
        total_questions += 1

    buffer = BytesIO()
    doc.save(buffer)

    st.success(f"{total_questions} questions detected.")

    st.download_button(
        "📥 Download Word",
        data=buffer.getvalue(),
        file_name="formatted_mcqs.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
